import sys
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from markdown_it import MarkdownIt
from mdit_py_plugins.texmath import texmath_plugin
from latex2mathml.converter import convert as latex_to_mathml


# ======================
# 基础格式
# ======================

def set_font(run, name="仿宋", size=16, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_line_spacing(p):
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(28)


def set_layout(doc):
    section = doc.sections[0]
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.4)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.3)


# ======================
# LaTeX 公式处理
# ======================

def insert_math(paragraph, latex):
    mathml = latex_to_mathml(latex)
    run = paragraph.add_run(f"[公式:{latex}]")
    set_font(run)


# ======================
# 主转换逻辑
# ======================

def convert(md_path, output):

    md = MarkdownIt().use(texmath_plugin)

    with open(md_path, "r", encoding="utf-8") as f:
        text = f.read()

    tokens = md.parse(text)

    doc = Document()
    set_layout(doc)

    i = 0
    while i < len(tokens):
        token = tokens[i]

        # ======================
        # 标题
        # ======================
        if token.type == "heading_open":
            level = int(token.tag[1])
            content = tokens[i+1].content

            p = doc.add_paragraph()

            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(content)
                set_font(run, "方正小标宋_GBK", 22, False)

            else:
                run = p.add_run(content)
                set_font(run, "仿宋", 16, False)
                set_line_spacing(p)

            i += 2

        # ======================
        # 段落
        # ======================
        elif token.type == "paragraph_open":
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)

            inline = tokens[i+1]
            for child in inline.children:

                if child.type == "text":
                    run = p.add_run(child.content)
                    set_font(run)

                elif child.type == "strong_open":
                    pass

                elif child.type == "strong_close":
                    pass

                elif child.type == "strong":
                    run = p.add_run(child.content)
                    set_font(run, bold=True)

                elif child.type == "math_inline":
                    insert_math(p, child.content)

                elif child.type == "softbreak":
                    p.add_run("\n")

            set_line_spacing(p)
            i += 2

        # ======================
        # 有序列表
        # ======================
        elif token.type == "ordered_list_open":
            i += 1
            while tokens[i].type != "ordered_list_close":
                if tokens[i].type == "inline":
                    p = doc.add_paragraph(style="List Number")
                    run = p.add_run(tokens[i].content)
                    set_font(run)
                i += 1

        # ======================
        # 无序列表
        # ======================
        elif token.type == "bullet_list_open":
            i += 1
            while tokens[i].type != "bullet_list_close":
                if tokens[i].type == "inline":
                    p = doc.add_paragraph(style="List Bullet")
                    run = p.add_run(tokens[i].content)
                    set_font(run)
                i += 1

        else:
            i += 1

    doc.save(output)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("python md_to_docx.py input.md output.docx")
        sys.exit(1)

    convert(sys.argv[1], sys.argv[2])